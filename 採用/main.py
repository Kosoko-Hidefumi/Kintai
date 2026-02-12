import os
import re
import pandas as pd
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from openpyxl import load_workbook
import shutil
from datetime import datetime

# Set up paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_DIR = os.path.join(BASE_DIR, 'input')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
TEMPLATE_DIR = os.path.join(BASE_DIR, 'template')
DATA_DIR = os.path.join(BASE_DIR, 'data')
APPLICANT_LIST_PATH = os.path.join(DATA_DIR, 'applicant_list.xlsx')

# Ensure directories exist
os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

# Path to Tesseract executable (Update this if necessary)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Path to Poppler bin directory
POPPLER_PATH = os.path.join(BASE_DIR, 'poppler-25.12.0', 'Library', 'bin')

# Path to Custom Tessdata (Local)
TESSDATA_DIR = os.path.join(BASE_DIR, 'tessdata')

import subprocess
import tempfile

def run_tesseract(image):
    """
    Runs tesseract on a PIL Image by saving to temp file and calling subprocess.
    Handles encoding issues (UTF-8 vs CP932).
    """
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp:
        image.save(temp.name)
        temp_path = temp.name
    
    try:
        cmd = [
            pytesseract.pytesseract.tesseract_cmd,
            temp_path,
            'stdout',
            '-l', 'jpn',
            '--tessdata-dir', TESSDATA_DIR
        ]
        
        # Run tesseract
        result = subprocess.run(cmd, capture_output=True)
        
        if result.returncode != 0:
            print(f"Tesseract Error: {result.stderr.decode('cp932', errors='replace')}")
            return ""

        # Try decoding stdout
        try:
            return result.stdout.decode('utf-8')
        except UnicodeDecodeError:
            return result.stdout.decode('cp932', errors='replace')
            
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file using OCR.
    """
    try:
        images = convert_from_path(pdf_path, poppler_path=POPPLER_PATH)
        text = ""
        
        for image in images:
            text += run_tesseract(image)
            
        return text
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error processing {pdf_path}: {e}")
        return None

def parse_applicant_info(text):
    """
    Parses extracted text to find Name, Zip, Address, Phone.
    """
    info = {
        '名前': None,
        '郵便番号': None,
        '住所': None,
        '電話番号': None
    }
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    # 1. Zip Code (Strongest signal)
    # Looking for 〒 or just pattern
    zip_match = re.search(r'(?:〒|郵便番号)?\s*(\d{3}[-]\d{4})', text)
    if zip_match:
        info['郵便番号'] = zip_match.group(1)

    # 2. Phone Number
    # Try labeled first
    phone_match = re.search(r'(?:電話(?:番号)?|TEL|携帯)[:：\s]*([\d-]{10,13})', text)
    if phone_match:
        info['電話番号'] = phone_match.group(1)
    else:
        # Fallback: Look for phone pattern (090..., 080..., 03... etc)
        # Exclude dates like 2026-02-10 (4digits-2digits-2digits)
        # Phone usually 2/3/4 digits - 3/4 digits - 4 digits
        # Regex: ^0\d{1,4}-\d{1,4}-\d{3,4}$ loosely
        dates = re.findall(r'\d{4}[-/年]', text) # Identify dates to avoid
        
        # Find all potential phone numbers
        candidates = re.findall(r'(0\d{1,4}-?\d{1,4}-?\d{3,4})', text)
        for c in candidates:
            # Clean up
            c_clean = c.replace('-', '')
            if len(c_clean) >= 10 and not any(c.startswith(d[:4]) for d in dates):
                # Prioritize mobile (090, 080, 070)
                if c_clean.startswith(('090', '080', '070')):
                    info['電話番号'] = c
                    break
                # Or landline
                if info['電話番号'] is None:
                    info['電話番号'] = c

    # 3. Address
    # Try labeled "住所" or "現住所"
    # Often the value is on the same line OR the NEXT line
    address_found = False
    for i, line in enumerate(lines):
        # Skip instruction lines (e.g. "現住所以外に連絡を希望する場合のみ記入")
        if '連絡を希望する' in line: continue

        if '住所' in line and (i+1 < len(lines)):
            # Check if address is on this line
            remainder = re.sub(r'.*住所[:：\s]*', '', line).strip()
            
            # If remainder contains instruction, clear it
            if '連絡を希望する' in remainder: remainder = ''

            # If remainder is just Zip code, look at next line
            if re.match(r'〒?\d{3}-\d{4}', remainder) or not remainder:
                candidate = lines[i+1]
                # Validate candidate (should contain prefecture or city or numbers)
                if any(x in candidate for x in ['県', '都', '府', '道', '市', '区']):
                    info['住所'] = candidate
                    address_found = True
                    break
            elif len(remainder) > 3: # Assuming valid address length
                 info['住所'] = remainder
                 address_found = True
                 break
    
    if not address_found:
        # heuristic: look for line starting with prefecture
        prefectures = ['北海道', '青森', '岩手', '宮城', '秋田', '山形', '福島', '茨城', '栃木', '群馬', '埼玉', '千葉', '東京', '神奈川', '新潟', '富山', '石川', '福井', '山梨', '長野', '岐阜', '静岡', '愛知', '三重', '滋賀', '京都', '大阪', '兵庫', '奈良', '和歌山', '鳥取', '島根', '岡山', '広島', '山口', '徳島', '香川', '愛媛', '高知', '福岡', '佐賀', '長崎', '熊本', '大分', '宮崎', '鹿児島', '沖縄']
        for line in lines:
            for p in prefectures:
                if line.startswith(p):
                    info['住所'] = line
                    break
            if info['住所']: break

    # 4. Name
    # Hardest one. Look for "氏名"
    # Or look for isolated line early in doc
    name_found = False
    
    # Strategy A: Label search
    name_match = re.search(r'(?:氏名|名前)[:：\s]+([^\s\n]+)', text)
    if name_match:
        info['名前'] = name_match.group(1)
        name_found = True
        
    # Strategy B: Heuristic (First non-keyword line, often large text in OCR, but we just have order)
    if not name_found or not info['名前']:
        # Skip lines that contain: 履歴書, 令和, 平成, 昭和, Date, Gender, ふりがな
        # Skip lines that are just numbers or single chars
        # Often Name is the first "Meaningful" line after Title
        for line in lines[:20]: # Check first 20 lines
            line = line.replace(' ', '').replace('　', '') # Remove spaces
            if '履歴書' in line: continue
            if '現在' in line: continue # Date line
            if 'ふりがな' in line: continue
            if 'フリガナ' in line: continue
            if '氏名' in line: continue # Label line
            if re.search(r'\d{4}年', line): continue
            if re.search(r'満\d+歳', line): continue
            
            # Candidate? Must be 2+ chars, mostly Kanji/Kana
            if 2 <= len(line) <= 10:
                # Assume this is name
                # Retrieve original line with spaces
                original_line = [l for l in lines if l.replace(' ', '').replace('　', '') == line][0]
                info['名前'] = original_line
                break
                
    return info

def update_excel(applicant_info):
    """
    Updates the Excel ledger with applicant info.
    Returns the updated row data including pass/fail status.
    """
    if not os.path.exists(APPLICANT_LIST_PATH):
        # Create new workbook if it doesn't exist (with Japanese headers)
        df = pd.DataFrame(columns=['名前', '郵便番号', '住所', '電話番号', '一次合否', '二次合否', '面接日', '面接時間'])
        df.to_excel(APPLICANT_LIST_PATH, index=False)
    
    # Read existing data
    df = pd.read_excel(APPLICANT_LIST_PATH)
    
    # Ensure necessary columns exist
    for col in ['一次合否', '二次合否', '面接日', '面接時間']:
        if col not in df.columns:
            df[col] = ''
    
    # Check if applicant already exists (by Name and Phone)
    existing = df[
        (df['名前'] == applicant_info['名前']) & 
        (df['電話番号'] == applicant_info['電話番号'])
    ]
    
    if existing.empty:
        # Append new applicant
        new_row = applicant_info.copy()
        new_row['一次合否'] = '' # Initialize as empty
        new_row['二次合否'] = ''
        # 面接日/時間は空のまま
        # new_row['面接日'] = '' # specific assignment not needed if empty
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        # Save back to Excel
        df.to_excel(APPLICANT_LIST_PATH, index=False)
        # Return the last row (which is the new one) as a Series/dict
        # However, newly added row has empty status.
        # We need to return it as dict.
        # Re-fetch or just construct.
        # Constructing is safer.
        return df.iloc[-1].to_dict()
    else:
        # Return existing data so we can check pass/fail status
        return existing.iloc[0].to_dict()

def generate_word_doc(template_name, context, output_filename):
    """
    Generates a Word document by replacing placeholders.
    """
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if value is None: value = ""
            if f'{{{{{key}}}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
    
    # Also check tables if templates use them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in context.items():
                         if value is None: value = ""
                         if f'{{{{{key}}}}}' in paragraph.text:
                            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"Generated: {output_path}")

def main():
    # Check if input directory has files
    files = [f for f in os.listdir(INPUT_DIR) if f.lower().endswith('.pdf')]
    if not files:
        print("No PDF files found in 'input' directory.")
        return

    for filename in files:
        pdf_path = os.path.join(INPUT_DIR, filename)
        print(f"Processing: {filename}")
        
        # 1. Extract Text
        text = extract_text_from_pdf(pdf_path)
        if not text:
            print(f"Failed to extract text from {filename}")
            continue
            
        print("Extracted Text Preview:", text[:100].replace('\n', ' '))
        
        # 2. Parse Info
        info = parse_applicant_info(text)
        print("Parsed Info:", info)
        
        if not info['名前']:
            print("Could not identify Name. Skipping.")
            continue

        # 3. Update Excel
        applicant_data = update_excel(info)
        
        # 4. Generate Documents based on status
        name = applicant_data.get('名前')
        first_pass = str(applicant_data.get('一次合否', '')).strip()
        second_pass = str(applicant_data.get('二次合否', '')).strip()
        
        # 1st Pass Document
        if first_pass in ['◯', '○', 'O', 'OK']:
            context = {
                'name': name,
                '面接日': applicant_data.get('面接日', ''),
                '面接時間': applicant_data.get('面接時間', '')
            }
            generate_word_doc('1st_goukaku.docx', context, f'{name}_1st_goukaku.docx')
            
        # 2nd Fail Document
        # If 1st passed AND 2nd NOT passed (and likely explicitly failed or process done?)
        # For safety, let's stick to simple logic: 1st=O and 2nd!=O
        if first_pass in ['◯', '○', 'O', 'OK']:
            if second_pass not in ['◯', '○', 'O', 'OK']:
                context = {
                    'name': name
                }
                generate_word_doc('2nd_fugoukaku.docx', context, f'{name}_2nd_fugoukaku.docx')

if __name__ == "__main__":
    main()