import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook

# Set up paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, 'template')
DATA_DIR = os.path.join(BASE_DIR, 'data')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output', '1st_pass')
APPLICANT_LIST_PATH = os.path.join(DATA_DIR, 'applicant_list.xlsx')

def generate_word_doc(template_name, context, output_filename):
    """
    Generates a Word document by replacing placeholders.
    """
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if value is None: value = ""
            # Try matching with parentheses first (Japanese and English)
            # User specifically asked for （面接日）, so we prioritize that if it exists.
            keys_to_try = [f'（{key}）', f'({key})', key]
            
            for k in keys_to_try:
                if k in paragraph.text:
                    paragraph.text = paragraph.text.replace(k, str(value))
                    # If we replaced the specific key, stop trying other variations for this key in this paragraph
                    # (though unlikely to overlap in a way that matters for simple replacements)
                    break
    
    # Also check tables if templates use them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in context.items():
                        if value is None: value = ""
                        # Try matching with parentheses first (Japanese and English)
                        keys_to_try = [f'（{key}）', f'({key})', key]
                        
                        for k in keys_to_try:
                            if k in paragraph.text:
                                paragraph.text = paragraph.text.replace(k, str(value))
                                break

    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    try:
        doc.save(output_path)
        print(f"Generated: {output_path}")
    except PermissionError:
        print(f"PermissionError: Could not save to {output_path}. File might be open.")
        # Try saving with underscore
        base, ext = os.path.splitext(output_filename)
        new_output_filename = f"{base}_new{ext}"
        new_output_path = os.path.join(OUTPUT_DIR, new_output_filename)
        try:
            doc.save(new_output_path)
            print(f"Generated (renamed): {new_output_path}")
        except Exception as e:
            print(f"Failed to save renamed file: {e}")
    except Exception as e:
        print(f"Error saving {output_path}: {e}")

def main():
    # Check if data file exists
    if not os.path.exists(APPLICANT_LIST_PATH):
        print(f"Data file not found: {APPLICANT_LIST_PATH}")
        return

    # Read data
    try:
        df = pd.read_excel(APPLICANT_LIST_PATH)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    # Check required columns
    required_cols = ['名前', '一次合否', '面接日', '面接時間']
    for col in required_cols:
        if col not in df.columns:
            print(f"Error: Column '{col}' not found in Excel file.")
            return

    # Filter for 1st pass applicants
    # Accepting various forms of "Circle" marks often used in Japan
    pass_marks = ['◯', '○', 'O', 'OK', 'ok', 'Pass', 'pass', '合格']
    
    # Ensure '一次合否' is string for comparison
    df['一次合否'] = df['一次合否'].astype(str)
    
    passed_applicants = df[df['一次合否'].isin(pass_marks)]

    if passed_applicants.empty:
        print("No applicants found with 1st pass status.")
        return

    print(f"Found {len(passed_applicants)} applicants passed 1st round.")

    for index, row in passed_applicants.iterrows():
        name = row['名前']
        
        # Format Date: YYYY年MM月DD日
        raw_date = row['面接日']
        date_str = ''
        if pd.notna(raw_date):
            try:
                # Check if it's numeric (Excel serial date)
                if isinstance(raw_date, (int, float)):
                    dt = pd.to_datetime(raw_date, unit='D', origin='1899-12-30')
                else:
                    # Force convert to datetime for strings/timestamps
                    dt = pd.to_datetime(raw_date)
                
                date_str = dt.strftime('%Y年%m月%d日')
            except:
                # Fallback to string if parsing fails
                date_str = str(raw_date)

        # Format Time: HH:MM
        raw_time = row['面接時間']
        time_str = ''
        if pd.notna(raw_time):
             # Try parsing as time or datetime
             try:
                 # If it's a string like '10:00:00' or '10:00'
                 s = str(raw_time)
                 # Simple regex check or just slice if it looks like time
                 if ':' in s:
                     parts = s.split(':')
                     if len(parts) >= 2:
                         time_str = f"{parts[0].zfill(2)}:{parts[1].zfill(2)}"
                     else:
                         time_str = s
                 else:
                     # Maybe it's a datetime/time object
                     if hasattr(raw_time, 'strftime'):
                         time_str = raw_time.strftime('%H:%M')
                     else:
                         time_str = s
             except:
                 time_str = str(raw_time)

        # Prepare context for template
        # User requested to insert into （面接日） and （面接時間）.
        # Since checking the template revealed NO parentheses around these words,
        # we will target the plain words '面接日' and '面接時間' as before,
        # BUT we must be careful not to match the label if it was "日時：".
        # The dump showed: "日時：　面接日　　　　　　面接時間"
        # So '面接日' acts as the placeholder.
        # We will apply the formatted strings.
        context = {
            '応募者名': name,
            '面接日': date_str,
            '面接時間': time_str
        }
        
        output_filename = f'{name}_1st_goukaku.docx'
        generate_word_doc('1st_goukaku.docx', context, output_filename)

if __name__ == "__main__":
    main()
