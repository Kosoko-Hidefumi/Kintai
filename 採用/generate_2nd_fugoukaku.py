import os
import pandas as pd
from docx import Document
import datetime

# Set up paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, 'template')
DATA_DIR = os.path.join(BASE_DIR, 'data')
# Output to output/2nd_fugoukaku
OUTPUT_DIR = os.path.join(BASE_DIR, 'output', '2nd_fugoukaku')
APPLICANT_LIST_PATH = os.path.join(DATA_DIR, 'applicant_list.xlsx')

def generate_word_doc(template_name, context, output_filename):
    """
    Generates a Word document by replacing placeholders.
    """
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    if not os.path.exists(template_path):
        print(f"Template not found: {template_path}")
        return

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if value is None: value = ""
            # Try matching with parentheses first (Japanese and English)
            keys_to_try = [f'（{key}）', f'({key})', key]
            
            for k in keys_to_try:
                if k in paragraph.text:
                    paragraph.text = paragraph.text.replace(k, str(value))
                    break
    
    # Also check tables if templates use them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in context.items():
                        if value is None: value = ""
                        keys_to_try = [f'（{key}）', f'({key})', key]
                        
                        for k in keys_to_try:
                            if k in paragraph.text:
                                paragraph.text = paragraph.text.replace(k, str(value))
                                break

    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    try:
        doc.save(output_path)
        print(f"Generated: {output_path}")
    except PermissionError:
        print(f"PermissionError: Could not save to {output_path}. File might be open.")
        # Try saving with underscore
        base, ext = os.path.splitext(output_filename)
        new_output_filename = f"{base}_new{ext}"
        new_output_path = os.path.join(OUTPUT_DIR, new_output_filename)
        try:
            doc.save(new_output_path)
            print(f"Generated (renamed): {new_output_path}")
        except Exception as e:
            print(f"Failed to save renamed file: {e}")
    except Exception as e:
        print(f"Error saving {output_path}: {e}")

def main():
    # Check if data file exists
    if not os.path.exists(APPLICANT_LIST_PATH):
        print(f"Data file not found: {APPLICANT_LIST_PATH}")
        return

    # Read data
    try:
        df = pd.read_excel(APPLICANT_LIST_PATH)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    # Check required columns
    required_cols = ['名前', '二次合否']
    for col in required_cols:
        if col not in df.columns:
            print(f"Error: Column '{col}' not found in Excel file.")
            return

    # Filter conditions:
    # 1. '二次合否' is empty (NaN or empty string)
    # Note: df['二次合否'].isna() covers NaN. 
    # Also treating empty string '' as empty.
    
    # Convert to string to handle various types, replace 'nan' with empty
    # fillna('') is safer before astype(str) for object columns containing NaNs
    df['二次合否'] = df['二次合否'].fillna('').astype(str)
    
    # Also handle string 'nan' if it somehow got in
    df['二次合否'] = df['二次合否'].replace('nan', '')
    
    # Filter: Second pass status is empty
    target_applicants = df[df['二次合否'] == '']

    if target_applicants.empty:
        print("No applicants found with empty 2nd pass status.")
        return

    print(f"Found {len(target_applicants)} applicants with empty 2nd pass status.")

    for index, row in target_applicants.iterrows():
        name = row['名前']
        
        # Prepare context for template
        context = {
            '応募者名': name
        }
        
        output_filename = f'{name}_2nd_fugoukaku.docx'
        generate_word_doc('2nd_fugoukaku.docx', context, output_filename)

if __name__ == "__main__":
    main()
