from docx import Document
import os

template_path = r'd:\code4biz\採用\template\1st_goukaku.docx'

if not os.path.exists(template_path):
    print("Template not found.")
else:
    print(f"Checking template: {template_path}")
    doc = Document(template_path)
    
    print("--- Paragraphs ---")
    for para in doc.paragraphs:
        if '{{' in para.text:
            print(f"Found placeholder in para: {para.text}")
    
    print("--- Tables ---")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{' in para.text:
                        print(f"Found placeholder in table: {para.text}")

    print("--- Headers ---")
    for section in doc.sections:
        for para in section.header.paragraphs:
            if '{{' in para.text:
                print(f"Found placeholder in header: {para.text}")
            
    print("--- Done ---")
