import os
import pandas as pd
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Set up paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
APPLICANT_LIST_PATH = os.path.join(DATA_DIR, 'applicant_list.xlsx')

# KOKUYO KPC-E121 specs (Approximate for Word)
# A4 Size
PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)

# Margins
MARGIN_TOP = Mm(15) 
MARGIN_BOTTOM = Mm(15)
MARGIN_LEFT = Mm(7.2)
MARGIN_RIGHT = Mm(7.2)

# Label Size
LABEL_WIDTH = Mm(63.5)
LABEL_HEIGHT = Mm(38.1)
HORIZONTAL_GAP = Mm(2.54)
VERTICAL_GAP = Mm(0) # Usually 0 for this type if they touch, or small if not. KPC-E121 usually has 0 vertical gap between rows.

def set_cell_margins(cell, top=0, start=100, bottom=0, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for m, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def create_label_content(cell, row_data):
    # Set internal cell margins to avoid text touching edges
    # set_cell_margins(cell, top=100, start=100, bottom=100, end=100) # excessive? check later.
    
    # 〒 Zip
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"〒 {row_data.get('郵便番号', '')}")
    run.font.size = Pt(10)
    
    # Address
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    addr = str(row_data.get('住所', ''))
    # Simple wrap if too long? Word handles wrapping within cell.
    run = p.add_run(addr)
    run.font.size = Pt(9)
    
    # Name
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10) # Add some space before name
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = str(row_data.get('名前', ''))
    run = p.add_run(f"{name}  様")
    run.font.size = Pt(14)
    run.bold = True

def main():
    if not os.path.exists(APPLICANT_LIST_PATH):
        print("Data file not found.")
        return

    df = pd.read_excel(APPLICANT_LIST_PATH)
    applicants = df.to_dict('records')
    print(f"Found {len(applicants)} applicants.")

    doc = Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)

    # Calculate rows per page (7 rows x 3 cols = 21 labels)
    items_per_page = 21
    
    # Chunk applicants
    for i in range(0, len(applicants), items_per_page):
        page_applicants = applicants[i:i + items_per_page]
        
        # Create table for this page
        # 5 columns: Label, Gap, Label, Gap, Label
        table = doc.add_table(rows=7, cols=5)
        table.autofit = False   
        table.allow_autofit = False

        # Set column widths
        # Col 0, 2, 4 are Labels (63.5mm)
        # Col 1, 3 are Gaps (2.54mm)
        widths = [LABEL_WIDTH, HORIZONTAL_GAP, LABEL_WIDTH, HORIZONTAL_GAP, LABEL_WIDTH]
        
        for row in table.rows:
            row.height = LABEL_HEIGHT
            row.height_rule = 2 # WD_ROW_HEIGHT_RULE.EXACTLY
            
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Fill data
        for j, applicant in enumerate(page_applicants):
            # Map index j (0-20) to table cell
            # Row = j // 3
            # Col = (j % 3) * 2  (0->0, 1->2, 2->4)
            r = j // 3
            c = (j % 3) * 2
            
            cell = table.cell(r, c)
            create_label_content(cell, applicant)
        
        # Add page break if there are more applicants
        if i + items_per_page < len(applicants):
            doc.add_page_break()

    output_path = os.path.join(OUTPUT_DIR, 'labels.docx')
    try:
        doc.save(output_path)
        print(f"Generated: {output_path}")
    except PermissionError:
        print(f"PermissionError: Could not save to {output_path}. File might be open.")
        output_path = os.path.join(OUTPUT_DIR, 'labels_new.docx')
        doc.save(output_path)
        print(f"Generated (renamed): {output_path}")

if __name__ == "__main__":
    main()
