from docx import Document
import os

template_path = r'd:\code4biz\採用\template\2nd_fugoukaku.docx'

if os.path.exists(template_path):
    print(f"Dumping text from: {template_path}")
    doc = Document(template_path)
    all_text = []
    
    for para in doc.paragraphs:
        all_text.append(para.text)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append(para.text)

    full_text = "\n".join(all_text)
    print(full_text)
else:
    print("Template not found.")
