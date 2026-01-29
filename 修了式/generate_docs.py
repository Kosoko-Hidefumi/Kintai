
import pandas as pd
from docx import Document
import os
import shutil

def generate_docs():
    data_file = "list.xlsx"
    output_dir = "output_docs"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    # Read Data
    try:
        df = pd.read_excel(data_file)
        # Clean columns
        df.columns = [c.strip() for c in df.columns]
        print(f"Columns: {df.columns.tolist()}")
    except Exception as e:
        print(f"Error reading {data_file}: {e}")
        return

    # Sort DataFrame: '祝辞' == '祝辞' comes first
    if '祝辞' in df.columns:
        # Create a temp column for sorting
        df['sort_key'] = df['祝辞'].apply(lambda x: 1 if str(x).strip() == '祝辞' else 0)
        df = df.sort_values(by='sort_key', ascending=False)
        # Drop temp column if desired, or just ignore
    
    # Templates
    template_general = "template_general.docx"
    template_guest = "template_guest.docx"
    
    count = 0
    ordered_pdfs = [] # Track order for merging
    
    for index, row in df.iterrows():
        name = str(row['名前']).strip()
        title = str(row['肩書き']).strip() if pd.notna(row['肩書き']) else ""
        shukuji = str(row['祝辞']).strip() if '祝辞' in df.columns and pd.notna(row['祝辞']) else ""
        
        # Determine Template
        template_path = template_guest if shukuji == '祝辞' else template_general
        
        if not os.path.exists(template_path):
            print(f"Template not found: {template_path}")
            continue
            
        try:
            doc = Document(template_path)
            
            # Replace logic
            # Search for "役職" and "名前" separately
            
            for p in doc.paragraphs:
                # Replace '役職'
                if '役職' in p.text:
                    p.text = p.text.replace('役職', title)

                # Replace '名前'
                if '名前' in p.text:
                    p.text = p.text.replace('名前', name)
            
            # Additional check for tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if '役職' in p.text:
                                p.text = p.text.replace('役職', title)
                            if '名前' in p.text:
                                p.text = p.text.replace('名前', name)
            
            # Save
            out_filename = f"案内状_{name}.docx"
            # Sanitize filename
            out_filename = "".join([c for c in out_filename if c.isalnum() or c in (' ', '.', '_', '(', ')', '-')])
            out_path = os.path.join(output_dir, out_filename)
            doc.save(out_path)
            print(f"Generated: {out_filename} (Template: {template_path})")
            
            # Add corresponding PDF name to strict order list
            ordered_pdfs.append(out_filename.replace('.docx', '.pdf'))
            count += 1
            
        except Exception as e:
            print(f"Error processing {name}: {e}")

    print(f"Generated {count} documents. Converting to PDF...")
    
    # PDF Conversion and Merging
    try:
        from docx2pdf import convert
        from pypdf import PdfWriter
        
        pdf_merger = PdfWriter()
        generated_pdfs = []
        
        print("Converting all docx to pdf...")
        convert(output_dir) # Batch convert matches filenames
        
        # Merge in the order of 'ordered_pdfs'
        print("Merging PDFs in sorted order (Guests first)...")
        for pdf_name in ordered_pdfs:
            pdf_path = os.path.join(output_dir, pdf_name)
            if os.path.exists(pdf_path):
                pdf_merger.append(pdf_path)
                generated_pdfs.append(pdf_path)
            else:
                print(f"Warning: PDF not found for merge: {pdf_name}")
            
        merged_output = "案内状_一括_v2.pdf"
        pdf_merger.write(merged_output)
        pdf_merger.close()
        
        print(f"Merged PDF created: {merged_output}")
        
        # Cleanup individual PDFs
        print("Cleaning up individual PDF files...")
        for pdf in generated_pdfs:
            try:
                os.remove(pdf)
            except Exception as e:
                print(f"Error removing {pdf}: {e}")
        
    except ImportError:
        print("Error: docx2pdf or pypdf not installed.")
    except Exception as e:
        print(f"Error converting/merging: {e}")

    print(f"Completed.")

if __name__ == "__main__":
    generate_docs()
