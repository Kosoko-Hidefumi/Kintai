
import pandas as pd
from docx import Document
from docx2pdf import convert
from pypdf import PdfWriter
import os
import shutil

def generate_orei():
    data_file = "list.xlsx"
    output_dir = "output_orei"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    else:
        # Clean directory to avoid mixing old run
        shutil.rmtree(output_dir)
        os.makedirs(output_dir)

    # Read Data
    try:
        df = pd.read_excel(data_file)
        # Clean columns
        df.columns = [c.strip() for c in df.columns]
        print(f"Columns: {df.columns.tolist()}")
    except Exception as e:
        print(f"Error reading {data_file}: {e}")
        return

    # Filter by Photo Distribution (写真配布)
    # Check what the mark is. Assuming non-null/non-empty is a mark, or specific char.
    if '写真配布' in df.columns:
        # Filter: keep row if '写真配布' is not NaN and contains '○' or '〇' or just is truthy?
        # Let's inspect unique values in logic or assume common marks.
        # User said "◯印" (Large circle? or Kanji zero?).
        target_marks = ['○', '〇', 'OK', '有', '◯']
        # Convert to string and check
        df = df[df['写真配布'].apply(lambda x: pd.notna(x) and str(x).strip() in target_marks)]
        print(f"Filtered target records: {len(df)}")
    else:
        print("Column '写真配布' not found. Cannot filter.")
        return

    # Sort DataFrame: '祝辞' == '祝辞' comes first for merging order
    if '祝辞' in df.columns:
        df['sort_key'] = df['祝辞'].apply(lambda x: 1 if str(x).strip() == '祝辞' else 0)
        df = df.sort_values(by='sort_key', ascending=False)
    
    # Templates
    template_shukuji = "orei_shukuji.docx"
    template_shukuden = "orei_shukuden.docx"
    template_general = "orei_general.docx"
    
    count = 0
    ordered_pdfs = []
    
    for index, row in df.iterrows():
        name = str(row['名前']).strip()
        # Strip existing '殿' or '様' if present to avoid double honorifics with template
        for suffix in [' 殿', '殿', ' 様', '様']:
            if name.endswith(suffix):
                name = name[:-len(suffix)]
                break
        
        title = str(row['肩書き']).strip() if pd.notna(row['肩書き']) else ""
        
        shukuji = str(row['祝辞']).strip() if '祝辞' in df.columns and pd.notna(row['祝辞']) else ""
        shukuden = str(row['祝電']).strip() if '祝電' in df.columns and pd.notna(row['祝電']) else ""
        
        # Debug print
        # print(f"Processing {name}: Shukuji='{shukuji}', Shukuden='{shukuden}'")

        # Determine Template and Suffix
        suffix = ""
        if shukuji == '祝辞':
            template_path = template_shukuji
            suffix = "_祝辞"
        elif shukuden in ['○', '〇', 'OK', '有', '◯']: 
             template_path = template_shukuden
             suffix = "_祝電"
        else:
             template_path = template_general
        
        if not os.path.exists(template_path):
            print(f"Template not found: {template_path}")
            continue
            
        try:
            doc = Document(template_path)
            
            # Replace logic
            for p in doc.paragraphs:
                # Try multiple keywords for title
                for term in ['役職', '肩書', '肩書き']:
                    if term in p.text:
                        p.text = p.text.replace(term, title)
                
                if '名前' in p.text:
                    p.text = p.text.replace('名前', name)
            
            for table in doc.tables:
                for row_t in table.rows:
                    for cell in row_t.cells:
                        for p in cell.paragraphs:
                            for term in ['役職', '肩書', '肩書き']:
                                if term in p.text:
                                    p.text = p.text.replace(term, title)
                            if '名前' in p.text:
                                p.text = p.text.replace('名前', name)
            
            # Save
            out_filename = f"御礼状_{name}{suffix}.docx"
            out_filename = "".join([c for c in out_filename if c.isalnum() or c in (' ', '.', '_', '(', ')', '-')])
            out_path = os.path.join(output_dir, out_filename)
            doc.save(out_path)
            print(f"Generated: {out_filename} (Template: {template_path})")
            
            ordered_pdfs.append(out_filename.replace('.docx', '.pdf'))
            count += 1
            
        except Exception as e:
            print(f"Error processing {name}: {e}")

    if count == 0:
        print("No documents generated. Check marks in '写真配布' column.")
        return

    print(f"Generated {count} documents. Converting to PDF...")
    
    try:
        print("Converting all docx to pdf...")
        convert(output_dir)
        
        print("Merging PDFs in sorted order...")
        pdf_merger = PdfWriter()
        generated_pdfs = []
        
        for pdf_name in ordered_pdfs:
            pdf_path = os.path.join(output_dir, pdf_name)
            if os.path.exists(pdf_path):
                pdf_merger.append(pdf_path)
                generated_pdfs.append(pdf_path)
            else:
                print(f"PDF missing: {pdf_name}")
                
        merged_output = "御礼状_一括.pdf"
        pdf_merger.write(merged_output)
        pdf_merger.close()
        print(f"Merged PDF created: {merged_output}")
        
        print("Cleaning up individual PDF files...")
        for pdf in generated_pdfs:
             try:
                os.remove(pdf)
             except:
                 pass
                 
    except Exception as e:
        print(f"Error converting/merging: {e}")

if __name__ == "__main__":
    generate_orei()
