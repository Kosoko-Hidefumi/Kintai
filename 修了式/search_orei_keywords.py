
from docx import Document
import os

files = ['orei_shukuji.docx', 'orei_shukuden.docx', 'orei_general.docx']
keywords = ['役職', '肩書', '肩書き', '名前']

print("Searching keywords...")
for f in files:
    if os.path.exists(f):
        print(f"--- {f} ---")
        try:
            doc = Document(f)
            found = False
            for i, p in enumerate(doc.paragraphs):
                for kw in keywords:
                    if kw in p.text:
                        print(f"FOUND '{kw}' in P{i}: '{p.text}'")
                        found = True
            
            if not found:
                 print("No keywords found in paragraphs. Checking tables...")
                 for t_idx, table in enumerate(doc.tables):
                     for r_idx, row in enumerate(table.rows):
                         for c_idx, cell in enumerate(row.cells):
                             for kw in keywords:
                                 if kw in cell.text:
                                      print(f"FOUND '{kw}' in Table{t_idx}R{r_idx}C{c_idx}: '{cell.text}'")
                                      found = True
            if not found:
                print("!! NO KEYWORDS FOUND !!")

        except Exception as e:
            print(f"Error: {e}")
