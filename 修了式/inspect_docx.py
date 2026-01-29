
from docx import Document
import os

files = ['template_general.docx', 'template_guest.docx']

for f in files:
    if os.path.exists(f):
        print(f"--- {f} ---")
        try:
            doc = Document(f)
            # Inspect first few paragraphs to find headers
            for i, p in enumerate(doc.paragraphs[:10]):
                print(f"P{i}: {p.text}")
                
            # Check tables if any
            for i, table in enumerate(doc.tables[:1]):
                 print(f"Table {i}:")
                 for row in table.rows:
                     print([cell.text for cell in row.cells])
        except Exception as e:
            print(f"Error reading {f}: {e}")
    else:
        print(f"{f} not found")
